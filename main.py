import argparse
import requests
import re
import urllib3
import openpyxl
import socket
import docx


from openpyxl.styles import *
from docx.text.paragraph import Paragraph
from docx.oxml.parser import OxmlElement


urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning) # enlève les msg de con de verify SSl
def enumeration_user(url):
   response = requests.get("https://"+url+"/wp-json/wp/v2/users", verify=False)
   if response.status_code == 200:
       users = response.json()
       print("\n📌 Utilisateurs WordPress trouvés :")
       for user in users:
           print(user["name"]+" :"+ user["slug"])
           return user["name"]+" :"+ user["slug"]
   else:
       print("Le fichier ", url+"/wp-json/wp/v2/users n'existe pas")


def version_wp(url):
  response = requests.get("https://"+url, verify=False)
  if re.search("WordPress [0-9.]{1,10}", response.text) is not None:
      wordpress_version = re.search("WordPress [0-9.]{1,10}", response.text).group()
      print(wordpress_version)
  else:
      print("unable to detect wordpress version")
      wordpress_version="Unknown"
  return wordpress_version


def detection_CMS(url):
   response = requests.get("https://" + url, verify=False)
   if re.search("WordPress [0-9.]{1,10}", response.text) is not None:
       wordpress_version = "yes"
   else:
       wordpress_version = "no"
   return wordpress_version

def get_ip_from_url(url):
   try:
       ip_address = socket.gethostbyname(url)
   except:
       ip_address = "Unknown"
   return ip_address

def insert_paragraph_after(paragraph, text=None, style=None):
   new_p = OxmlElement("w:p")
   paragraph._p.addnext(new_p)
   new_para = Paragraph(new_p, paragraph._parent)
   if text:
       new_para.add_run(text)
   elif style is not None:
       new_para.style = style
   return new_para


#permet de lancer le script py dans le terminal avec des options
parser = argparse.ArgumentParser(description="Donne des information sur les site wordpress")
parser.add_argument('-t', '--target', type=str, help="url du site sans de http")
parser.add_argument('-ue', '--user-enumeration', action='store_true', help="Enumere les utilisateurs")
parser.add_argument('-wv', '--wordpress-version', action='store_true', help="Donne la version de wordpress")
parser.add_argument('-if', '--input-file', type=str, help="Recupere des url dans un fichier")
parser.add_argument('-do', '--document-output', type=str, help="Chemin du rapport Word à compléter")




args=parser.parse_args()


if args.target:
   url = args.target
   users, version, ip = [], "Unknown", get_ip_from_url(url)
   if args.user_enumeration:
       enumeration_user(url)
   elif args.wordpress_version:
       version_wp(url)
   if args.document_output:
        document_path = args.document_output
        document = docx.Document(document_path)
        is_wordpress = detection_CMS(url)
        print(f"WordPress détecté : {is_wordpress}")
        
        for para in document.paragraphs:
            if para.style.name == "Heading 1":
                if para.text == "Adresse IP":
                    insert_paragraph_after(para, f"L'adresse IP du site ({url}) est {ip}")
                if para.text == "CMS":
                    if is_wordpress == "yes":
                        insert_paragraph_after(para, f"Le site {url} fonctionne sous WordPress. Version : {version}")
                        if users:
                            insert_paragraph_after(para, "Utilisateurs détectés : " + ", ".join(users))
        
        document.save(document_path)
        print(f"Rapport mis à jour avec succès : {document_path}")
   else:
       usr=enumeration_user(url)
       ver=version_wp(url)
       ip=get_ip_from_url(url)



elif args.input_file:
   path="Classeur1.xlsx"
   wb_obj =openpyxl.load_workbook(path)
   worksheet_1=wb_obj.active
   worksheet_1.column_dimensions["A"].width = 50
   worksheet_1.column_dimensions["C"].width = 20
   worksheet_1.column_dimensions["E"].width = 50
   print(worksheet_1["A1"].value)


   input_file=args.input_file


   with open(input_file, "r") as list_file:
       url_list=list_file.readlines()


   i=2
   for url in url_list:
       url = url.rstrip()
       ip_address= get_ip_from_url(url)
       print(url+": "+ip_address)
       worksheet_1["A"+str(i)] = url


       if ip_address == "Unknown":
           worksheet_1["B" + str(i)] = "No"
           worksheet_1["B" + str(i)].font = Font(color="FF0000")
           worksheet_1["E" + str(i)] = "Unknown"
       else:
           worksheet_1["B" + str(i)] = "Yes"
           worksheet_1["B" + str(i)].font = Font(color="00561b")
           resultat_version = version_wp(url)
           worksheet_1["E" + str(i)] = resultat_version
       worksheet_1["C"+str(i)] = ip_address
       i+=1
   wb_obj.save("Classeur1.xlsx")
else:
   print("error, no url specify, please use -t <url>")
   exit()
